#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
deliver.py — video-watch 一键生成阅读交付件（三大件中的两个）

读 run 目录下的 transcript.txt 与 frames/frames.json，生成：
  - 【文字稿】<标题>.docx（python-docx：标题/信息段/逐行转写，时间戳加粗，微软雅黑 10.5pt）
  - 【关键帧】<标题>.pdf （reportlab：封面页 + 每页 2 帧，帧图按原始宽高比缩放，
    宽约 16.5cm，图注 `帧 NNN / 总数 ｜ t = HH:MM:SS`，页眉标题、页脚页码）

文件名净化 _sanitize_name：去掉半角 /\\:*?"<>| 与控制字符、首尾 trim、最长 60 字符。
时间显示 _fmt_hms 遵循仓库惯例（common.fmt_ts）：MM:SS，满 1 小时用 HH:MM:SS。

CJK 字体多级回退（reportlab TTFont；.ttc 需 subfontIndex=0）：
  --font → 环境变量 DAIMON_CJK_FONT_REGULAR / DAIMON_CJK_FONT_BOLD
  → C:\\Windows\\Fonts\\msyh.ttc → macOS/Linux 常见 Noto CJK 路径
  → 全部失败则按契约 die 并提示用 --font 指定。

reportlab / python-docx 在函数内部惰性 import；模块顶层仅依赖 stdlib，
测试可无第三方包直接 import 纯辅助函数（_sanitize_name / _fmt_hms）。

参数：
  --run <run_dir>   必填，run 目录（含 transcript.txt 与 frames/frames.json）
  --out-dir <目录>  交付件输出目录（默认 = run 目录根部）
  --title <标题>    覆盖标题（默认 manifest.json 的 title，读不到用 run 目录名）
  --font <路径>     指定 CJK 字体文件（跳过自动回退探测）

stdout 最后一行：RESULT_JSON: {"ok": true, "title": ..., "outputs": {"docx": ..., "pdf": ...}, ...}
"""
from __future__ import annotations

import argparse
import json
import os
import re
import sys
from datetime import datetime
from pathlib import Path

SKILL_ROOT = Path(__file__).resolve().parent.parent  # skill 根目录
sys.path.insert(0, str(Path(__file__).resolve().parent))  # 任意 cwd 下可 import common

# ---------------------------------------------------------------------------
# 优先复用同目录 common.py；导入失败时启用最小兜底实现，保证脚本被
# 单独分发时仍可运行（与仓库其他脚本同款有意保留的兜底块）。
# ---------------------------------------------------------------------------
try:
    import common  # type: ignore

    fmt_ts = common.fmt_ts
    redact_url = common.redact_url
except Exception:  # pragma: no cover - common 缺失时的最小兜底实现

    def fmt_ts(sec):
        """时间戳显示：MM:SS，≥1 小时用 HH:MM:SS。"""
        sec = max(0.0, float(sec))
        h = int(sec // 3600)
        m = int((sec % 3600) // 60)
        s = int(round(sec % 60))
        if s == 60:
            s, m = 0, m + 1
        if m == 60:
            m, h = 0, h + 1
        return f"{h:02d}:{m:02d}:{s:02d}" if h else f"{m:02d}:{s:02d}"

    def redact_url(value):
        """兜底：不做 URL 凭据清理，原样返回。"""
        return str(value)


# ---------------------------------------------------------------------------
# 输出与错误处理（契约：日志在前，stdout 最后一行单行 RESULT_JSON）
# ---------------------------------------------------------------------------
def log(msg):
    print(f"[deliver] {msg}", flush=True)


def emit_result(obj):
    print("RESULT_JSON: " + json.dumps(obj, ensure_ascii=False, default=str), flush=True)


def fail(msg, code=1):
    print(f"[deliver][ERROR] {msg}", file=sys.stderr, flush=True)
    emit_result({"ok": False, "error": str(msg)})
    sys.exit(code)


class CliParser(argparse.ArgumentParser):
    """参数错误也遵守 RESULT_JSON 契约。"""

    def error(self, message):
        self.print_usage(sys.stderr)
        fail(f"参数错误: {message}")


# ---------------------------------------------------------------------------
# 纯 stdlib 辅助函数（模块顶层可直接 import，无第三方依赖）
# ---------------------------------------------------------------------------

# Windows 文件名半角非法字符 + 所有控制字符（C0/C1/DEL）
_ILLEGAL_NAME_RE = re.compile(r'[/\\:*?"<>|\x00-\x1f\x7f-\x9f]')

MAX_NAME_LEN = 60


def _sanitize_name(text, max_len: int = MAX_NAME_LEN, fallback: str = "video") -> str:
    """净化标题为安全文件名片段。

    去掉半角 /\\:*?"<>| 与控制字符，首尾 trim，最长 max_len 字符；
    净化结果为空时回退 fallback。
    """
    cleaned = _ILLEGAL_NAME_RE.sub("", str(text or ""))
    cleaned = cleaned.strip()[:max(1, int(max_len))].rstrip()
    return cleaned or fallback


def _fmt_hms(seconds) -> str:
    """时间显示，遵循仓库惯例：MM:SS，满 1 小时用 HH:MM:SS（秒先四舍五入）。"""
    try:
        return fmt_ts(float(seconds))
    except (TypeError, ValueError):
        return "00:00"


# 转写行格式：[MM:SS] 文本 或 [HH:MM:SS] 文本
_TRANSCRIPT_LINE_RE = re.compile(r"^\[(\d{1,3}:\d{2}(?::\d{2})?)\]\s?(.*)$")

# CJK 字体自动探测候选（--font 与环境变量优先；.ttc 一律 subfontIndex=0）
_FONT_CANDIDATES_REGULAR = (
    r"C:\Windows\Fonts\msyh.ttc",
    "/System/Library/Fonts/PingFang.ttc",
    "/Library/Fonts/Arial Unicode.ttf",
    "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
    "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
    "/usr/share/fonts/noto-cjk/NotoSansCJK-Regular.ttc",
)
_FONT_CANDIDATES_BOLD = (
    r"C:\Windows\Fonts\msyhbd.ttc",
    "/usr/share/fonts/opentype/noto/NotoSansCJK-Bold.ttc",
    "/usr/share/fonts/truetype/noto/NotoSansCJK-Bold.ttc",
    "/usr/share/fonts/noto-cjk/NotoSansCJK-Bold.ttc",
)


# ---------------------------------------------------------------------------
# 输入读取
# ---------------------------------------------------------------------------

def _load_manifest(run_dir: Path) -> dict:
    """尽力读取 manifest.json；缺失/损坏返回 {}（标题等字段有兜底，不视为致命）。"""
    path = run_dir / "manifest.json"
    try:
        data = json.loads(path.read_text(encoding="utf-8-sig"))
        return data if isinstance(data, dict) else {}
    except Exception:
        return {}


def _resolve_title(run_dir: Path, manifest: dict, override: str | None) -> str:
    """标题来源：--title 覆盖 → manifest.title → run 目录名；统一经 _sanitize_name 净化。"""
    raw = (override or "").strip() or str(manifest.get("title") or "").strip() or run_dir.name
    return _sanitize_name(raw, fallback="video")


def _load_transcript(run_dir: Path) -> list[str]:
    """读取 transcript.txt 的非空行（保留原文，含 [MM:SS] 前缀）。"""
    path = run_dir / "transcript.txt"
    if not path.is_file():
        fail(f"转写文件不存在: {path}")
    try:
        lines = path.read_text(encoding="utf-8-sig").splitlines()
    except OSError as exc:
        fail(f"无法读取转写文件 {path}: {exc}")
    return [ln.rstrip() for ln in lines if ln.strip()]


def _load_frames(run_dir: Path) -> list[dict]:
    """读取 frames/frames.json（顶层为帧对象列表），按时间排序返回。"""
    path = run_dir / "frames" / "frames.json"
    if not path.is_file():
        fail(f"帧索引不存在: {path}")
    try:
        data = json.loads(path.read_text(encoding="utf-8-sig"))
    except (OSError, json.JSONDecodeError) as exc:
        fail(f"无法解析帧索引 {path}: {exc}")
    if isinstance(data, dict):  # 容忍包装形式 {"frames": [...]}
        data = data.get("frames") or data.get("entries") or []
    if not isinstance(data, list):
        fail(f"帧索引格式不支持（应为列表）: {path}")
    frames = [f for f in data if isinstance(f, dict) and f.get("file")]
    frames.sort(key=lambda f: _frame_time(f))
    return frames


def _frame_time(frame: dict) -> float:
    """帧时间优先级：actual_t → requested_t → t。"""
    for key in ("actual_t", "requested_t", "t"):
        try:
            value = frame.get(key)
            if value is not None:
                return float(value)
        except (TypeError, ValueError):
            continue
    return 0.0


# ---------------------------------------------------------------------------
# DOCX（python-docx，惰性 import）
# ---------------------------------------------------------------------------

def _build_docx(
    out_path: Path,
    *,
    title: str,
    manifest: dict,
    transcript_lines: list[str],
    frame_count: int,
) -> None:
    from docx import Document  # 惰性 import python-docx
    from docx.oxml.ns import qn
    from docx.shared import Pt

    CJK_FONT = "微软雅黑"

    def _style_run(run, *, bold: bool = False, size: float = 10.5) -> None:
        run.font.name = CJK_FONT
        run.font.size = Pt(size)
        run.bold = bold
        # 中文字体必须同时设置 eastAsia，否则 Word 里中文回落默认字体
        run._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)  # noqa: SLF001

    doc = Document()
    heading = doc.add_heading("", level=1)
    _style_run(heading.add_run(title), bold=True, size=16)

    info_bits = []
    source = manifest.get("input") or manifest.get("video_path") or ""
    if source:
        info_bits.append(f"来源：{redact_url(source)}")
    duration = manifest.get("duration")
    if isinstance(duration, (int, float)) and duration > 0:
        info_bits.append(f"时长：{_fmt_hms(duration)}")
    info_bits.append(f"转写行数：{len(transcript_lines)}")
    info_bits.append(f"关键帧数：{frame_count}")
    info_bits.append(f"生成日期：{datetime.now().strftime('%Y-%m-%d')}")
    info_para = doc.add_paragraph()
    _style_run(info_para.add_run(" ｜ ".join(info_bits)), size=9)

    for line in transcript_lines:
        para = doc.add_paragraph()
        match = _TRANSCRIPT_LINE_RE.match(line)
        if match:
            _style_run(para.add_run(f"[{match.group(1)}] "), bold=True)
            _style_run(para.add_run(match.group(2)))
        else:
            _style_run(para.add_run(line))

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


# ---------------------------------------------------------------------------
# PDF（reportlab，惰性 import）
# ---------------------------------------------------------------------------

def _resolve_cjk_fonts(font_override: str | None):
    """按回退链注册 CJK 字体，返回 (regular_name, bold_name)。全部失败 → fail()。"""
    from reportlab.pdfbase import pdfmetrics  # 惰性 import reportlab
    from reportlab.pdfbase.ttfonts import TTFont

    def _try_register(name: str, path: str) -> bool:
        try:
            kwargs = {"subfontIndex": 0} if path.lower().endswith((".ttc", ".otc")) else {}
            pdfmetrics.registerFont(TTFont(name, path, **kwargs))
            return True
        except Exception:
            return False

    def _resolve(candidates, env_var: str, label: str) -> str | None:
        ordered = []
        env_path = os.environ.get(env_var, "").strip()
        if env_path:
            ordered.append(env_path)
        ordered.extend(candidates)
        for path in ordered:
            if path and Path(path).is_file() and _try_register(f"VWCJK-{label}", path):
                log(f"CJK {label} 字体: {path}")
                return f"VWCJK-{label}"
        return None

    if font_override:
        if not Path(font_override).is_file():
            fail(f"--font 指定的字体文件不存在: {font_override}")
        if not _try_register("VWCJK-regular", font_override):
            fail(f"--font 指定的字体无法被 reportlab 加载: {font_override}")
        log(f"CJK 字体（--font 指定）: {font_override}")
        return "VWCJK-regular", "VWCJK-regular"

    regular = _resolve(_FONT_CANDIDATES_REGULAR, "DAIMON_CJK_FONT_REGULAR", "regular")
    if regular is None:
        fail(
            "未找到可用 CJK 字体，PDF 中文将无法渲染。请用 --font <字体路径> 指定，"
            "或设置环境变量 DAIMON_CJK_FONT_REGULAR（粗体用 DAIMON_CJK_FONT_BOLD）。"
        )
    bold = _resolve(_FONT_CANDIDATES_BOLD, "DAIMON_CJK_FONT_BOLD", "bold")
    if bold is None:
        log("未找到独立粗体 CJK 字体，粗体复用常规字体")
        bold = regular
    return regular, bold


def _build_pdf(
    out_path: Path,
    *,
    title: str,
    manifest: dict,
    frames: list[dict],
    frames_dir: Path,
    transcript_count: int,
    font_override: str | None,
) -> None:
    from reportlab.lib.pagesizes import A4  # 惰性 import reportlab
    from reportlab.lib.units import cm
    from reportlab.lib.utils import ImageReader
    from reportlab.pdfgen import canvas as pdf_canvas

    regular, bold = _resolve_cjk_fonts(font_override)

    PAGE_W, PAGE_H = A4
    IMG_W = 16.5 * cm
    MARGIN_X = (PAGE_W - IMG_W) / 2
    CAPTION_H = 0.7 * cm
    GAP = 0.5 * cm
    TOP_Y = PAGE_H - 2.2 * cm  # 页眉下方起始

    total = len(frames)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    c = pdf_canvas.Canvas(str(out_path), pagesize=A4)

    def _header_footer(page_num: int, *, with_header: bool = True) -> None:
        if with_header:
            c.setFont(regular, 9)
            c.drawCentredString(PAGE_W / 2, PAGE_H - 1.2 * cm, title)
            c.setLineWidth(0.5)
            c.line(MARGIN_X, PAGE_H - 1.5 * cm, PAGE_W - MARGIN_X, PAGE_H - 1.5 * cm)
        c.setFont(regular, 9)
        c.drawCentredString(PAGE_W / 2, 1.0 * cm, f"第 {page_num} 页")

    # ---- 封面页 ----
    page_num = 1
    c.setFont(bold, 20)
    c.drawCentredString(PAGE_W / 2, PAGE_H - 6 * cm, "视频关键帧图集")
    c.setFont(bold, 14)
    # 标题过长时截断显示（文件名已限 60 字符，封面再保险截到 40）
    c.drawCentredString(PAGE_W / 2, PAGE_H - 7.2 * cm, title[:40])
    c.setFont(regular, 11)
    cover_lines = []
    source = manifest.get("input") or manifest.get("video_path") or ""
    if source:
        cover_lines.append(f"来源：{redact_url(source)}")
    duration = manifest.get("duration")
    if isinstance(duration, (int, float)) and duration > 0:
        cover_lines.append(f"时长：{_fmt_hms(duration)}")
    created = str(manifest.get("created_at") or "").strip()
    if created:
        cover_lines.append(f"阅读时间：{created}")
    cover_lines.append(f"关键帧数：{total}")
    cover_lines.append(f"转写行数：{transcript_count}")
    cover_lines.append(f"生成日期：{datetime.now().strftime('%Y-%m-%d')}")
    y = PAGE_H - 9.5 * cm
    for line in cover_lines:
        c.drawString(MARGIN_X, y, line)
        y -= 0.9 * cm
    _header_footer(page_num, with_header=False)
    c.showPage()

    # ---- 正文：每页 2 帧 ----
    idx = 0
    while idx < total:
        page_num += 1
        y = TOP_Y
        for _ in range(2):
            if idx >= total:
                break
            frame = frames[idx]
            img_path = frames_dir / str(frame["file"])
            n = idx + 1
            t_disp = _fmt_hms(_frame_time(frame))
            caption = f"帧 {n:03d} / {total} ｜ t = {t_disp}"
            img_h = 0.0
            if img_path.is_file():
                try:
                    reader = ImageReader(str(img_path))
                    w, h = reader.getSize()
                    img_h = IMG_W * (h / w) if w else 0.0
                    if img_h > 0:
                        c.drawImage(
                            reader,
                            MARGIN_X,
                            y - img_h,
                            width=IMG_W,
                            height=img_h,
                            preserveAspectRatio=True,
                            mask="auto",
                        )
                except Exception as exc:
                    log(f"跳过无法读取的帧图 {img_path.name}: {exc}")
            else:
                log(f"帧图缺失，仅输出图注: {img_path}")
            c.setFont(regular, 10)
            c.drawCentredString(PAGE_W / 2, y - img_h - 0.5 * cm, caption)
            y = y - img_h - CAPTION_H - GAP
            idx += 1
        _header_footer(page_num)
        c.showPage()

    c.save()


# ---------------------------------------------------------------------------
# main
# ---------------------------------------------------------------------------

def main(argv=None) -> int:
    for stream in (sys.stdout, sys.stderr):
        try:
            stream.reconfigure(encoding="utf-8", errors="replace")  # type: ignore[attr-defined]
        except Exception:
            pass

    ap = CliParser(
        prog="deliver.py",
        description="video-watch 交付件生成：读 run 目录的 transcript.txt 与 frames/frames.json，"
                    "生成 【文字稿】<标题>.docx 与 【关键帧】<标题>.pdf。",
        epilog="RESULT_JSON 字段: ok, title, outputs{docx,pdf}, stats{frames,transcript_lines}, out_dir。",
    )
    ap.add_argument("--run", dest="run_dir", required=True, help="run 目录（必填）")
    ap.add_argument("--out-dir", default=None, help="交付件输出目录（默认 = run 目录根部）")
    ap.add_argument("--title", default=None, help="覆盖标题（默认 manifest.json 的 title，兜底 run 目录名）")
    ap.add_argument("--font", default=None, help="CJK 字体文件路径（跳过自动回退探测）")
    args = ap.parse_args(argv)

    run_dir = Path(args.run_dir)
    if not run_dir.is_dir():
        fail(f"run 目录不存在: {run_dir}")
    out_dir = Path(args.out_dir) if args.out_dir else run_dir

    try:
        manifest = _load_manifest(run_dir)
        title = _resolve_title(run_dir, manifest, args.title)
        log(f"标题: {title}")

        transcript_lines = _load_transcript(run_dir)
        frames = _load_frames(run_dir)
        log(f"输入: 转写 {len(transcript_lines)} 行，关键帧 {len(frames)} 帧")

        docx_path = out_dir / f"【文字稿】{title}.docx"
        pdf_path = out_dir / f"【关键帧】{title}.pdf"

        _build_docx(
            docx_path,
            title=title,
            manifest=manifest,
            transcript_lines=transcript_lines,
            frame_count=len(frames),
        )
        log(f"DOCX 已生成 -> {docx_path}")

        _build_pdf(
            pdf_path,
            title=title,
            manifest=manifest,
            frames=frames,
            frames_dir=run_dir / "frames",
            transcript_count=len(transcript_lines),
            font_override=args.font,
        )
        log(f"PDF 已生成 -> {pdf_path}")

        emit_result({
            "ok": True,
            "title": title,
            "outputs": {"docx": str(docx_path.resolve()), "pdf": str(pdf_path.resolve())},
            "stats": {"frames": len(frames), "transcript_lines": len(transcript_lines)},
            "out_dir": str(out_dir.resolve()),
        })
        return 0
    except SystemExit:
        raise
    except Exception as e:
        fail(f"未预期错误: {type(e).__name__}: {e}")


if __name__ == "__main__":
    sys.exit(main())
